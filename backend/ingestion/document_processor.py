"""
Document Intelligence Agent - Ingestion Core
---------------------------------------------
Multi-format document processor. Converts any supported industrial
document into normalized plain text + structured entities, ready for
chunking, embedding, and knowledge-graph construction.

Supported formats: PDF (native + scanned/OCR), DOCX, XLSX/CSV, images
(JPEG/PNG incl. P&ID screenshots and equipment nameplate photos), TXT/JSON.
"""
from __future__ import annotations

import json
import csv
import io
from pathlib import Path
from dataclasses import dataclass

import fitz  # PyMuPDF
import docx
import openpyxl

from backend.ingestion.ocr import ocr_image_bytes
from backend.ingestion.extractors import extract_entities, ExtractedEntities
from backend.config import CHUNK_SIZE, CHUNK_OVERLAP

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf", ".docx": "docx", ".doc": "docx",
    ".xlsx": "xlsx", ".xls": "xlsx", ".csv": "csv",
    ".png": "image", ".jpg": "image", ".jpeg": "image",
    ".txt": "text", ".json": "json", ".xml": "xml",
}


@dataclass
class IngestedDocument:
    filename: str
    doc_type: str
    raw_text: str
    entities: ExtractedEntities
    category: str


def infer_category(filename: str, text: str) -> str:
    name = filename.lower()
    low = text.lower()[:3000]
    if any(k in name or k in low for k in ["work order", "wo-", "maintenance"]):
        return "Maintenance Report"
    if any(k in name or k in low for k in ["inspection"]):
        return "Inspection Report"
    if any(k in name or k in low for k in ["sop", "procedure"]):
        return "SOP / Procedure"
    if any(k in name or k in low for k in ["incident", "accident"]):
        return "Incident Report"
    if any(k in name or k in low for k in ["iso", "compliance", "audit", "peso", "oisd", "factory act"]):
        return "Compliance Document"
    if any(k in name or k in low for k in ["p&id", "pid", "drawing", "schematic"]):
        return "Engineering Drawing"
    if any(k in name or k in low for k in ["manual"]):
        return "Manual"
    return "General Document"


def _extract_pdf(path: Path) -> str:
    text_parts = []
    doc = fitz.open(path)
    for page in doc:
        page_text = page.get_text()
        if page_text.strip():
            text_parts.append(page_text)
        else:
            # Likely a scanned page -> rasterize + OCR
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            ocr_text = ocr_image_bytes(img_bytes)
            text_parts.append(ocr_text)
    doc.close()
    return "\n".join(text_parts)


def _extract_docx(path: Path) -> str:
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"--- Sheet: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(c) for c in row if c is not None]
            if row_vals:
                parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    parts = []
    with open(path, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            parts.append(" | ".join(row))
    return "\n".join(parts)


def _extract_image(path: Path) -> str:
    with open(path, "rb") as f:
        return ocr_image_bytes(f.read())


def _extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    return json.dumps(data, indent=2)


EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
    "csv": _extract_csv,
    "image": _extract_image,
    "text": _extract_text,
    "json": _extract_json,
    "xml": _extract_text,
}


def process_document(path: Path) -> IngestedDocument:
    ext = path.suffix.lower()
    doc_type = SUPPORTED_EXTENSIONS.get(ext)
    if not doc_type:
        raise ValueError(f"Unsupported file type: {ext}")

    extractor = EXTRACTORS[doc_type]
    raw_text = extractor(path) or ""
    entities = extract_entities(raw_text)
    category = infer_category(path.name, raw_text)

    return IngestedDocument(
        filename=path.name,
        doc_type=doc_type,
        raw_text=raw_text,
        entities=entities,
        category=category,
    )


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    """Semantic-ish sliding window chunking with parent-child structure.
    Returns list of (child_text, parent_text) tuples: each child is a small
    window used for precise retrieval, paired with a larger parent window
    used for grounded answer generation (parent-child chunking)."""
    words = text.split()
    if not words:
        return []

    chunks = []
    step = max(chunk_size - overlap, 1)
    parent_size = chunk_size * 3

    i = 0
    while i < len(words):
        child = " ".join(words[i:i + chunk_size])
        parent_start = max(0, i - chunk_size)
        parent_end = min(len(words), i + chunk_size + parent_size)
        parent = " ".join(words[parent_start:parent_end])
        if child.strip():
            chunks.append((child, parent))
        i += step
    return chunks
